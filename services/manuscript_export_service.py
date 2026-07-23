import io
import re
from collections import defaultdict
from html import escape
from pathlib import Path

from PIL import Image as PillowImage

from services.manuscript_asset_service import (
    read_manuscript_asset_file,
    render_equation_png,
)


ASSET_TYPES = ("figure", "table", "equation")
ASSET_TOKEN_PATTERN = re.compile(r"\[\[(figure|table|equation):(\d+)\]\]")
INLINE_MARKDOWN_PATTERN = re.compile(
    r"(?P<bolditalic>\*\*\*(?P<bolditalic_text>.+?)\*\*\*)"
    r"|(?P<bold>\*\*(?P<bold_text>.+?)\*\*)"
    r"|(?P<italic_star>(?<!\*)\*(?P<italic_star_text>[^*\n]+?)\*(?!\*))"
    r"|(?P<italic_under>(?<!\w)_(?P<italic_under_text>[^_\n]+?)_(?!\w))"
    r"|(?P<link>\[(?P<link_text>[^\]]+)\]\((?P<link_url>(?:https?://|mailto:)[^)\s]+)\))"
    r"|(?P<code>`(?P<code_text>[^`\n]+)`)"
)


def _author_parts(authors: str) -> list[str]:
    return [part.strip() for part in re.split(r"[;]", authors or "") if part.strip()]


def _author_surname(author: str) -> str:
    if "," in author:
        return author.split(",", 1)[0].strip() or "Unknown"

    words = re.findall(r"[\wÀ-ž'-]+", author)
    return words[-1] if words else "Unknown"


def _apa_inline(source) -> str:
    authors = _author_parts(source["authors"] or "")
    year = source["publication_year"] or "n.d."

    if not authors:
        author_label = source["title"] or "Unknown source"
    elif len(authors) == 1:
        author_label = _author_surname(authors[0])
    elif len(authors) == 2:
        author_label = f"{_author_surname(authors[0])} & {_author_surname(authors[1])}"
    else:
        author_label = f"{_author_surname(authors[0])} et al."

    return f"({author_label}, {year})"


def _apa_reference(source) -> str:
    authors = source["authors"] or "Unknown author"
    year = source["publication_year"] or "n.d."
    title = source["title"] or "Untitled source"
    journal = source["source_name"] or ""
    doi = source["doi"] or ""
    url = source["url"] or ""
    tail = f" https://doi.org/{doi}" if doi else (f" {url}" if url else "")
    journal_part = f" *{journal}*." if journal else ""
    return f"{authors} ({year}). {title}.{journal_part}{tail}".strip()


def _vancouver_reference(source, number: int) -> str:
    authors = source["authors"] or "Unknown author"
    title = source["title"] or "Untitled source"
    journal = source["source_name"] or ""
    year = source["publication_year"] or "n.d."
    doi = source["doi"] or ""
    tail = f" doi:{doi}." if doi else ""
    journal_part = f" {journal}." if journal else ""
    return f"{number}. {authors}. {title}.{journal_part} {year}.{tail}".strip()


def render_citations(text: str, sources, style: str) -> str:
    by_key = {source["citation_key"].casefold(): source for source in sources}
    vancouver_numbers = {
        source["citation_key"].casefold(): index
        for index, source in enumerate(sources, start=1)
    }

    def replace(match):
        key = match.group(1).strip().casefold()
        source = by_key.get(key)

        if not source:
            return match.group(0)

        if style == "Vancouver":
            return f"[{vancouver_numbers[key]}]"

        return _apa_inline(source)

    return re.sub(r"\[@([^\]]+)\]", replace, text or "")


def bibliography_lines(sources, style: str) -> list[str]:
    ordered = list(sources)

    if style == "APA 7":
        ordered.sort(key=lambda source: (
            str(source["authors"] or "").casefold(),
            str(source["title"] or "").casefold(),
        ))
        return [_apa_reference(source) for source in ordered]

    return [
        _vancouver_reference(source, index)
        for index, source in enumerate(ordered, start=1)
    ]


def _normalized_assets(assets) -> list[dict]:
    counters = {asset_type: 0 for asset_type in ASSET_TYPES}
    normalized = []

    for raw_asset in assets or []:
        asset = dict(raw_asset)
        asset_type = asset.get("asset_type")

        if asset_type not in ASSET_TYPES:
            continue

        counters[asset_type] += 1
        asset.setdefault("number", counters[asset_type])
        asset.setdefault("label", f"{asset_type.title()} {asset['number']}")
        content = asset.get("content", {})

        if not isinstance(content, dict):
            content = {}

        asset["content"] = content
        normalized.append(asset)

    return normalized


def render_asset_references(text: str, assets) -> str:
    labels = {
        (asset["asset_type"], int(asset["id"])): asset["label"]
        for asset in _normalized_assets(assets)
    }

    def replace(match):
        return labels.get((match.group(1), int(match.group(2))), match.group(0))

    return ASSET_TOKEN_PATTERN.sub(replace, text or "")


def render_manuscript_text(text: str, sources, style: str, assets) -> str:
    return render_asset_references(render_citations(text, sources, style), assets)


def _inline_markdown_parts(text: str) -> list[dict]:
    parts = []
    position = 0

    for match in INLINE_MARKDOWN_PATTERN.finditer(str(text or "")):
        if match.start() > position:
            parts.append({"kind": "text", "text": text[position:match.start()]})

        if match.group("bolditalic"):
            parts.append({
                "kind": "bolditalic",
                "text": match.group("bolditalic_text"),
            })
        elif match.group("bold"):
            parts.append({"kind": "bold", "text": match.group("bold_text")})
        elif match.group("italic_star"):
            parts.append({
                "kind": "italic",
                "text": match.group("italic_star_text"),
            })
        elif match.group("italic_under"):
            parts.append({
                "kind": "italic",
                "text": match.group("italic_under_text"),
            })
        elif match.group("link"):
            parts.append({
                "kind": "link",
                "text": match.group("link_text"),
                "url": match.group("link_url"),
            })
        else:
            parts.append({"kind": "code", "text": match.group("code_text")})

        position = match.end()

    if position < len(text):
        parts.append({"kind": "text", "text": text[position:]})

    return parts or [{"kind": "text", "text": str(text or "")}]


def _markdown_blocks(text: str) -> list[dict]:
    blocks = []
    paragraph_lines = []

    def flush_paragraph():
        if paragraph_lines:
            blocks.append({
                "kind": "paragraph",
                "text": " ".join(line.strip() for line in paragraph_lines).strip(),
            })
            paragraph_lines.clear()

    for raw_line in str(text or "").splitlines():
        line = raw_line.rstrip()

        if not line.strip():
            flush_paragraph()
            continue

        heading = re.match(r"^\s*(#{1,6})\s+(.+)$", line)
        bullet = re.match(r"^(?P<indent>\s*)[-*+]\s+(?P<text>.+)$", line)
        numbered = re.match(
            r"^(?P<indent>\s*)(?P<number>\d+)\.\s+(?P<text>.+)$",
            line,
        )
        quote = re.match(r"^\s*>\s?(?P<text>.+)$", line)

        if heading:
            flush_paragraph()
            blocks.append({
                "kind": "heading",
                "level": min(len(heading.group(1)), 3),
                "text": heading.group(2).strip(),
            })
        elif bullet:
            flush_paragraph()
            blocks.append({
                "kind": "bullet",
                "level": len(bullet.group("indent").replace("\t", "    ")) // 4,
                "text": bullet.group("text").strip(),
            })
        elif numbered:
            flush_paragraph()
            blocks.append({
                "kind": "number",
                "level": len(numbered.group("indent").replace("\t", "    ")) // 4,
                "number": int(numbered.group("number")),
                "text": numbered.group("text").strip(),
            })
        elif quote:
            flush_paragraph()
            blocks.append({"kind": "quote", "text": quote.group("text").strip()})
        else:
            paragraph_lines.append(line)

    flush_paragraph()
    return blocks


def _inline_markdown_to_reportlab(text: str) -> str:
    output = []

    for part in _inline_markdown_parts(text):
        safe_text = escape(part["text"])

        if part["kind"] == "bolditalic":
            output.append(f"<b><i>{safe_text}</i></b>")
        elif part["kind"] == "bold":
            output.append(f"<b>{safe_text}</b>")
        elif part["kind"] == "italic":
            output.append(f"<i>{safe_text}</i>")
        elif part["kind"] == "code":
            output.append(f'<font name="Courier">{safe_text}</font>')
        elif part["kind"] == "link":
            output.append(
                f'<a href="{escape(part["url"], quote=True)}" color="#0d65d9">'
                f"<u>{safe_text}</u></a>"
            )
        else:
            output.append(safe_text)

    return "".join(output)


def _assets_by_section(assets) -> dict[int, list[dict]]:
    grouped = defaultdict(list)

    for asset in _normalized_assets(assets):
        grouped[int(asset["section_id"])].append(asset)

    return grouped


def _table_parts(asset) -> tuple[list[str], list[list[str]]]:
    content = asset.get("content", {})
    columns = [str(column) for column in content.get("columns", [])]
    raw_rows = content.get("rows", [])
    rows = []

    for raw_row in raw_rows if isinstance(raw_rows, list) else []:
        if isinstance(raw_row, dict):
            rows.append([str(raw_row.get(column, "") or "") for column in columns])
        elif isinstance(raw_row, list):
            rows.append([
                str(raw_row[index] if index < len(raw_row) else "")
                for index in range(len(columns))
            ])

    return columns, rows


def _markdown_asset(asset) -> str:
    label_caption = f"{asset['label']}. {asset.get('caption', '')}".strip()

    if asset["asset_type"] == "figure":
        alt_text = asset.get("alt_text") or asset.get("caption") or asset["label"]
        path = asset.get("storage_path") or "missing-figure"
        return f"![{alt_text}]({path})\n\n*{label_caption}*"

    if asset["asset_type"] == "equation":
        latex = asset.get("content", {}).get("latex", "")
        return f"$$\n{latex}\n$$\n\n*{label_caption}*"

    columns, rows = _table_parts(asset)

    if not columns:
        table_markdown = "_Empty table_"
    else:
        clean = lambda value: str(value).replace("|", "\\|").replace("\n", " ")
        lines = [
            "| " + " | ".join(clean(column) for column in columns) + " |",
            "| " + " | ".join("---" for _ in columns) + " |",
        ]
        lines.extend(
            "| " + " | ".join(clean(value) for value in row) + " |"
            for row in rows
        )
        table_markdown = "\n".join(lines)

    return f"*{label_caption}*\n\n{table_markdown}"


def manuscript_markdown(manuscript, sections, sources, assets=()) -> str:
    style = manuscript["citation_style"]
    normalized_assets = _normalized_assets(assets)
    grouped_assets = _assets_by_section(normalized_assets)
    parts = [f"# {manuscript['title']}"]
    references_written = False

    for section in sections:
        title = section["title"]
        parts.append(f"\n## {title}\n")

        if section["section_type"] == "references":
            lines = bibliography_lines(sources, style)
            parts.append("\n\n".join(lines) if lines else "No sources attached.")
            references_written = True
        else:
            parts.append(render_manuscript_text(
                section["content_md"],
                sources,
                style,
                normalized_assets,
            ))

        for asset in grouped_assets.get(int(section["id"]), []):
            parts.append(f"\n{_markdown_asset(asset)}\n")

    if sources and not references_written:
        parts.append("\n## References\n")
        parts.append("\n\n".join(bibliography_lines(sources, style)))

    return "\n".join(parts).strip() + "\n"


def _configure_docx_table(table, columns: int):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    table.autofit = False
    table_width = 9360
    column_width = max(1, table_width // max(columns, 1))
    table_properties = table._tbl.tblPr
    table_width_element = table_properties.first_child_found_in("w:tblW")

    if table_width_element is None:
        table_width_element = OxmlElement("w:tblW")
        table_properties.append(table_width_element)

    table_width_element.set(qn("w:w"), str(table_width))
    table_width_element.set(qn("w:type"), "dxa")
    indentation = OxmlElement("w:tblInd")
    indentation.set(qn("w:w"), "120")
    indentation.set(qn("w:type"), "dxa")
    table_properties.append(indentation)
    margins = OxmlElement("w:tblCellMar")

    for side, width in (("top", 80), ("left", 120), ("bottom", 80), ("right", 120)):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(width))
        node.set(qn("w:type"), "dxa")
        margins.append(node)

    table_properties.append(margins)
    grid = table._tbl.tblGrid

    for child in list(grid):
        grid.remove(child)

    for _ in range(columns):
        grid_column = OxmlElement("w:gridCol")
        grid_column.set(qn("w:w"), str(column_width))
        grid.append(grid_column)

    for row in table.rows:
        for cell in row.cells:
            width_element = cell._tc.tcPr.first_child_found_in("w:tcW")

            if width_element is None:
                width_element = OxmlElement("w:tcW")
                cell._tc.tcPr.append(width_element)

            width_element.set(qn("w:w"), str(column_width))
            width_element.set(qn("w:type"), "dxa")

    header_properties = table.rows[0]._tr.get_or_add_trPr()
    header_repeat = OxmlElement("w:tblHeader")
    header_repeat.set(qn("w:val"), "true")
    header_properties.append(header_repeat)


def _docx_caption(document, text: str, *, keep_with_next: bool = False):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    paragraph = document.add_paragraph(style="Caption")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = keep_with_next
    paragraph.paragraph_format.keep_together = True
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(55, 65, 81)
    run.italic = True
    return paragraph


def _docx_image(document, image_bytes: bytes, *, max_width=6.1, max_height=7.2):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches

    with PillowImage.open(io.BytesIO(image_bytes)) as image:
        width_px, height_px = image.size
    aspect = width_px / max(height_px, 1)
    width = min(max_width, max_height * aspect)
    height = width / max(aspect, 0.01)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.keep_together = True
    paragraph.add_run().add_picture(
        io.BytesIO(image_bytes),
        width=Inches(width),
        height=Inches(height),
    )


def _add_docx_asset(document, asset):
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    caption = f"{asset['label']}. {asset.get('caption', '')}".strip()

    if asset["asset_type"] == "figure":
        try:
            _docx_image(document, read_manuscript_asset_file(asset["storage_path"]))
        except (FileNotFoundError, ValueError, OSError):
            paragraph = document.add_paragraph("[Figure file unavailable]")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _docx_caption(document, caption)
        return

    if asset["asset_type"] == "equation":
        latex = asset.get("content", {}).get("latex", "")

        try:
            _docx_image(
                document,
                render_equation_png(latex),
                max_width=5.5,
                max_height=0.8,
            )
        except (ValueError, RuntimeError):
            paragraph = document.add_paragraph(latex)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _docx_caption(document, caption)
        return

    columns, rows = _table_parts(asset)
    _docx_caption(document, caption, keep_with_next=True)

    if not columns:
        document.add_paragraph("[Empty table]")
        return

    table = document.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"

    for index, column in enumerate(columns):
        cell = table.rows[0].cells[index]
        cell.text = column
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "E8F1FD")
        cell._tc.get_or_add_tcPr().append(shading)

        for run in cell.paragraphs[0].runs:
            run.bold = True

    for row in rows:
        cells = table.add_row().cells

        for index, value in enumerate(row):
            cells[index].text = value
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    _configure_docx_table(table, len(columns))


def _add_docx_hyperlink(paragraph, text: str, url: str):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.opc.constants import RELATIONSHIP_TYPE

    relationship_id = paragraph.part.relate_to(
        url,
        RELATIONSHIP_TYPE.HYPERLINK,
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0D65D9")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.extend((color, underline))
    run.append(run_properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _add_docx_inline(paragraph, text: str):
    from docx.shared import Pt, RGBColor

    for part in _inline_markdown_parts(text):
        if part["kind"] == "link":
            _add_docx_hyperlink(paragraph, part["text"], part["url"])
            continue

        run = paragraph.add_run(part["text"])

        if part["kind"] in ("bold", "bolditalic"):
            run.bold = True

        if part["kind"] in ("italic", "bolditalic"):
            run.italic = True

        if part["kind"] == "code":
            run.font.name = "Courier New"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(52, 64, 84)


def _add_docx_markdown_content(document, content: str):
    from docx.shared import Inches

    for block in _markdown_blocks(content):
        kind = block["kind"]

        if kind == "heading":
            paragraph = document.add_heading("", level=block["level"])
        elif kind == "bullet":
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.left_indent = Inches(
                0.375 + 0.25 * block.get("level", 0)
            )
            paragraph.paragraph_format.first_line_indent = Inches(-0.194)
        elif kind == "number":
            paragraph = document.add_paragraph(style="List Number")
            paragraph.paragraph_format.left_indent = Inches(
                0.375 + 0.25 * block.get("level", 0)
            )
            paragraph.paragraph_format.first_line_indent = Inches(-0.194)
        elif kind == "quote":
            quote_style = (
                "Intense Quote"
                if "Intense Quote" in [style.name for style in document.styles]
                else None
            )
            paragraph = document.add_paragraph(style=quote_style)
            paragraph.paragraph_format.left_indent = Inches(0.3)
            paragraph.paragraph_format.right_indent = Inches(0.3)
        else:
            paragraph = document.add_paragraph()

        _add_docx_inline(paragraph, block["text"])


def _set_docx_style_font(style, font_name: str, size, color=None):
    from docx.oxml.ns import qn

    style.font.name = font_name
    style.font.size = size
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font_name)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font_name)

    if color is not None:
        style.font.color.rgb = color


def _configure_submission_docx(document):
    from docx.enum.section import WD_ORIENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    page = document.sections[0]
    page.page_width = Inches(8.5)
    page.page_height = Inches(11)
    page.orientation = WD_ORIENT.PORTRAIT
    page.top_margin = Inches(1)
    page.bottom_margin = Inches(1)
    page.left_margin = Inches(1)
    page.right_margin = Inches(1)
    page.header_distance = Inches(0.492)
    page.footer_distance = Inches(0.492)
    normal = document.styles["Normal"]
    _set_docx_style_font(normal, "Calibri", Pt(11))
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    title_style = document.styles["Title"]
    _set_docx_style_font(title_style, "Calibri", Pt(20), RGBColor(11, 37, 69))
    title_style.font.bold = True
    title_style.paragraph_format.space_before = Pt(0)
    title_style.paragraph_format.space_after = Pt(8)

    heading_tokens = {
        "Heading 1": (16, "2E74B5", 18, 10),
        "Heading 2": (13, "2E74B5", 12, 6),
        "Heading 3": (12, "1F4D78", 8, 4),
    }

    for style_name, (size, color, before, after) in heading_tokens.items():
        style = document.styles[style_name]
        _set_docx_style_font(
            style,
            "Calibri",
            Pt(size),
            RGBColor.from_string(color),
        )
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = document.styles[style_name]
        _set_docx_style_font(style, "Calibri", Pt(11))
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208

    footer = page.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def _clean_profile_value(value) -> str:
    text = str(value or "").strip()
    return "" if text.casefold() in {"nan", "none"} else text


def _add_submission_title_block(document, manuscript, submission_profile):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    profile = submission_profile or {}
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)
    title_run = title.add_run(manuscript["title"])
    title_run.bold = True
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = RGBColor(11, 37, 69)
    authors = profile.get("authors") or []

    if authors:
        author_parts = []

        for author in authors:
            name = _clean_profile_value(author.get("name"))
            affiliations = _clean_profile_value(author.get("affiliations"))

            if name:
                author_parts.append(f"{name} [{affiliations}]" if affiliations else name)

        if author_parts:
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(5)
            run = paragraph.add_run(", ".join(author_parts))
            run.bold = True
            run.font.name = "Calibri"
            run.font.size = Pt(11)

    for affiliation in profile.get("affiliations") or []:
        label = _clean_profile_value(affiliation.get("label"))
        institution = _clean_profile_value(affiliation.get("institution"))
        location = _clean_profile_value(affiliation.get("location"))

        if not institution:
            continue

        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(2)
        text = f"[{label}] {institution}" if label else institution

        if location:
            text += f", {location}"

        run = paragraph.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(71, 84, 103)

    corresponding = profile.get("corresponding_author") or {}
    corresponding_name = _clean_profile_value(corresponding.get("name"))
    corresponding_email = _clean_profile_value(corresponding.get("email"))

    if corresponding_name or corresponding_email:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(
            "Corresponding author: "
            + " · ".join(value for value in (corresponding_name, corresponding_email) if value)
        )
        run.font.name = "Calibri"
        run.font.size = Pt(9.5)

    target_journal = _clean_profile_value(profile.get("target_journal"))

    if target_journal:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(10)
        run = paragraph.add_run(f"Prepared for submission to {target_journal}")
        run.italic = True
        run.font.name = "Calibri"
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(71, 84, 103)


def _add_docx_bibliography_entry(document, text: str):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    _add_docx_inline(paragraph, text)


def manuscript_docx(
    manuscript,
    sections,
    sources,
    assets=(),
    submission_profile=None,
) -> bytes:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("DOCX export requires the python-docx dependency.") from exc

    document = Document()
    document.core_properties.title = manuscript["title"]
    document.core_properties.subject = "Scientific manuscript"
    _configure_submission_docx(document)
    _add_submission_title_block(document, manuscript, submission_profile)
    corresponding = (submission_profile or {}).get("corresponding_author") or {}

    if corresponding.get("name"):
        document.core_properties.author = _clean_profile_value(corresponding["name"])

    style = manuscript["citation_style"]
    normalized_assets = _normalized_assets(assets)
    grouped_assets = _assets_by_section(normalized_assets)
    references_written = False

    for section in sections:
        document.add_heading(section["title"], level=1)

        if section["section_type"] == "references":
            lines = bibliography_lines(sources, style)

            for line in lines or ["No sources attached."]:
                _add_docx_bibliography_entry(document, line)

            references_written = True
        else:
            content = render_manuscript_text(
                section["content_md"],
                sources,
                style,
                normalized_assets,
            )

            _add_docx_markdown_content(document, content)

            if section["section_type"] == "abstract" and (submission_profile or {}).get("keywords"):
                paragraph = document.add_paragraph()
                label = paragraph.add_run("Keywords: ")
                label.bold = True
                paragraph.add_run(", ".join(submission_profile["keywords"]))

        for asset in grouped_assets.get(int(section["id"]), []):
            _add_docx_asset(document, asset)

    if sources and not references_written:
        document.add_heading("References", level=1)

        for line in bibliography_lines(sources, style):
            _add_docx_bibliography_entry(document, line)

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _register_pdf_font() -> str:
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    candidates = (
        Path("/System/Library/Fonts/Supplemental/Arial Unicode.ttf"),
        Path("/System/Library/Fonts/Supplemental/Arial.ttf"),
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
    )

    for path in candidates:
        if path.is_file():
            pdfmetrics.registerFont(TTFont("ResearchJournalUnicode", str(path)))
            return "ResearchJournalUnicode"

    return "Helvetica"


def _pdf_image(image_bytes: bytes, *, max_width: float, max_height: float):
    from reportlab.platypus import Image

    with PillowImage.open(io.BytesIO(image_bytes)) as image:
        width_px, height_px = image.size
    scale = min(max_width / max(width_px, 1), max_height / max(height_px, 1))
    return Image(
        io.BytesIO(image_bytes),
        width=max(1, width_px * scale),
        height=max(1, height_px * scale),
        hAlign="CENTER",
    )


def _pdf_asset_flowables(asset, styles, font_name, available_width):
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import KeepTogether, Paragraph, Spacer, Table, TableStyle

    caption_style = ParagraphStyle(
        "ResearchCaption",
        parent=styles["BodyText"],
        fontName=font_name,
        fontSize=8.5,
        leading=11,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#374151"),
        spaceBefore=3,
        spaceAfter=8,
    )
    caption = Paragraph(
        escape(f"{asset['label']}. {asset.get('caption', '')}"),
        caption_style,
    )

    if asset["asset_type"] == "figure":
        try:
            image = _pdf_image(
                read_manuscript_asset_file(asset["storage_path"]),
                max_width=available_width,
                max_height=125 * mm,
            )
            return [KeepTogether([image, caption])]
        except (FileNotFoundError, ValueError, OSError):
            return [Paragraph("[Figure file unavailable]", caption_style), caption]

    if asset["asset_type"] == "equation":
        latex = asset.get("content", {}).get("latex", "")

        try:
            equation = _pdf_image(
                render_equation_png(latex),
                max_width=available_width * 0.8,
                max_height=16 * mm,
            )
            return [KeepTogether([equation, caption])]
        except (ValueError, RuntimeError):
            return [KeepTogether([
                Paragraph(escape(latex), caption_style),
                caption,
            ])]

    columns, rows = _table_parts(asset)

    if not columns:
        return [caption, Paragraph("[Empty table]", caption_style)]

    cell_style = ParagraphStyle(
        "ResearchTableCell",
        parent=styles["BodyText"],
        fontName=font_name,
        fontSize=8,
        leading=10,
    )
    data = [
        [Paragraph(f"<b>{escape(column)}</b>", cell_style) for column in columns]
    ]
    data.extend([
        [Paragraph(escape(value), cell_style) for value in row]
        for row in rows
    ])
    table = Table(
        data,
        colWidths=[available_width / len(columns)] * len(columns),
        repeatRows=1,
        hAlign="CENTER",
    )
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8F1FD")),
        ("TEXTCOLOR", (0, 0), (-1, -1), colors.HexColor("#111827")),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#9CA3AF")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    return [caption, table, Spacer(1, 3 * mm)]


def _pdf_markdown_flowables(content: str, styles, font_name):
    from reportlab.lib import colors
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import Paragraph

    body_style = ParagraphStyle(
        "ResearchMarkdownBody",
        parent=styles["BodyText"],
        fontName=font_name,
        fontSize=10.5,
        leading=15,
        spaceAfter=8,
    )
    heading_two_style = ParagraphStyle(
        "ResearchMarkdownHeading2",
        parent=body_style,
        fontName=font_name,
        fontSize=12.5,
        leading=16,
        textColor=colors.HexColor("#1F4F85"),
        spaceBefore=7,
        spaceAfter=5,
    )
    heading_three_style = ParagraphStyle(
        "ResearchMarkdownHeading3",
        parent=body_style,
        fontName=font_name,
        fontSize=11.5,
        leading=15,
        textColor=colors.HexColor("#344054"),
        spaceBefore=5,
        spaceAfter=4,
    )
    quote_style = ParagraphStyle(
        "ResearchMarkdownQuote",
        parent=body_style,
        leftIndent=14,
        rightIndent=10,
        borderColor=colors.HexColor("#CBD5E1"),
        borderWidth=0.6,
        borderPadding=7,
        backColor=colors.HexColor("#F8FAFC"),
        textColor=colors.HexColor("#475467"),
    )
    flowables = []

    for block in _markdown_blocks(content):
        markup = _inline_markdown_to_reportlab(block["text"]) or "&#160;"
        kind = block["kind"]

        if kind == "heading":
            paragraph_style = (
                heading_three_style
                if block.get("level", 2) >= 3
                else heading_two_style
            )
            flowables.append(Paragraph(markup, paragraph_style))
        elif kind == "bullet":
            list_style = ParagraphStyle(
                f"ResearchBullet{block.get('level', 0)}",
                parent=body_style,
                leftIndent=16 + 14 * block.get("level", 0),
                firstLineIndent=-9,
                spaceAfter=3,
            )
            flowables.append(Paragraph(markup, list_style, bulletText="•"))
        elif kind == "number":
            list_style = ParagraphStyle(
                f"ResearchNumber{block.get('level', 0)}",
                parent=body_style,
                leftIndent=18 + 14 * block.get("level", 0),
                firstLineIndent=-13,
                spaceAfter=3,
            )
            flowables.append(Paragraph(
                markup,
                list_style,
                bulletText=f"{block.get('number', 1)}.",
            ))
        elif kind == "quote":
            flowables.append(Paragraph(markup, quote_style))
        else:
            flowables.append(Paragraph(markup, body_style))

    return flowables or [Paragraph("&#160;", body_style)]


def manuscript_pdf(manuscript, sections, sources, assets=()) -> bytes:
    try:
        from reportlab.lib.enums import TA_CENTER
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import mm
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
    except ImportError as exc:
        raise RuntimeError("PDF export requires the reportlab dependency.") from exc

    output = io.BytesIO()
    font_name = _register_pdf_font()
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ResearchTitle",
        parent=styles["Title"],
        fontName=font_name,
        fontSize=19,
        leading=23,
        alignment=TA_CENTER,
        spaceAfter=14,
    )
    heading_style = ParagraphStyle(
        "ResearchHeading",
        parent=styles["Heading1"],
        fontName=font_name,
        fontSize=14,
        leading=18,
        spaceBefore=10,
        spaceAfter=7,
    )
    body_style = ParagraphStyle(
        "ResearchBody",
        parent=styles["BodyText"],
        fontName=font_name,
        fontSize=10.5,
        leading=15,
        spaceAfter=8,
    )
    document = SimpleDocTemplate(
        output,
        pagesize=A4,
        rightMargin=22 * mm,
        leftMargin=22 * mm,
        topMargin=20 * mm,
        bottomMargin=20 * mm,
        title=manuscript["title"],
    )
    story = [Paragraph(escape(manuscript["title"]), title_style)]
    style = manuscript["citation_style"]
    normalized_assets = _normalized_assets(assets)
    grouped_assets = _assets_by_section(normalized_assets)
    references_written = False

    for section in sections:
        story.append(Paragraph(escape(section["title"]), heading_style))

        if section["section_type"] == "references":
            blocks = bibliography_lines(sources, style) or ["No sources attached."]
            references_written = True
        else:
            content = render_manuscript_text(
                section["content_md"],
                sources,
                style,
                normalized_assets,
            )
            blocks = None

        if blocks is not None:
            for block in blocks:
                story.append(Paragraph(escape(block) or "&#160;", body_style))
        else:
            story.extend(_pdf_markdown_flowables(content, styles, font_name))

        for asset in grouped_assets.get(int(section["id"]), []):
            story.extend(_pdf_asset_flowables(
                asset,
                styles,
                font_name,
                document.width,
            ))

    if sources and not references_written:
        story.append(Paragraph("References", heading_style))

        for line in bibliography_lines(sources, style):
            story.append(Paragraph(escape(line), body_style))

    story.append(Spacer(1, 5 * mm))
    document.build(story)
    return output.getvalue()
