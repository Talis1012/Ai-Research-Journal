import os
import tempfile
import unittest
import zipfile
from io import BytesIO
from pathlib import Path

from PIL import Image

from db.database import init_db
from db.library_queries import create_library_item
from db.queries import create_project
from db.writing_queries import (
    add_manuscript_version_comment,
    add_manuscript_section,
    attach_manuscript_evidence,
    attach_manuscript_source,
    create_manuscript_asset,
    create_manuscript,
    create_manuscript_version,
    delete_manuscript_section,
    duplicate_manuscript_version,
    get_manuscript_assets,
    get_manuscript,
    get_manuscript_ai_context,
    get_manuscript_evidence,
    get_manuscript_section,
    get_manuscript_sections,
    get_manuscript_sources,
    get_manuscript_submission_profile,
    get_manuscript_version,
    get_manuscript_version_comments,
    get_manuscript_versions,
    get_project_library_sources,
    insert_section_citation,
    insert_section_citations,
    insert_manuscript_asset_reference,
    move_manuscript_asset,
    restore_manuscript_version,
    restore_manuscript_section_from_version,
    update_manuscript_section,
    update_manuscript_ai_context,
    update_manuscript_asset,
    update_manuscript_source,
    update_manuscript_submission_profile,
    update_manuscript_version,
    validate_section_citations,
)
from services.manuscript_export_service import (
    manuscript_docx,
    manuscript_markdown,
    manuscript_pdf,
    render_asset_references,
)
from services.manuscript_asset_service import (
    read_manuscript_asset_file,
    save_figure_upload,
)
from services.writing_service import generate_writing_suggestion
from services.manuscript_review_service import (
    publication_readiness,
    run_manuscript_checks,
)


class FixedWritingProvider:
    def __init__(self):
        self.prompt = ""

    def generate_json(self, prompt):
        self.prompt = prompt
        return {
            "suggested_text": "Supported result [@smith2025].",
            "explanation": "Grounded in the attached abstract.",
            "evidence_used": [
                {
                    "source_type": "library",
                    "source_id": 1,
                    "label": "Attached paper",
                    "support": "Reports the selected outcome.",
                }
            ],
            "claims": [
                {
                    "claim": "Supported result",
                    "status": "supported",
                    "reason": "The abstract reports it.",
                    "citation_keys": ["smith2025"],
                }
            ],
        }


class WritingTestCase(unittest.TestCase):
    def setUp(self):
        self.temp_dir = tempfile.TemporaryDirectory()
        self.previous_db_path = os.environ.get("DATABASE_PATH")
        self.previous_asset_path = os.environ.get("MANUSCRIPT_ASSET_STORAGE_PATH")
        os.environ["DATABASE_PATH"] = str(Path(self.temp_dir.name) / "writing.db")
        os.environ["MANUSCRIPT_ASSET_STORAGE_PATH"] = str(
            Path(self.temp_dir.name) / "manuscript-assets"
        )
        init_db()
        self.project_id = create_project(
            "Writing project",
            "Medicinal chemistry",
            "Evidence-grounded manuscript",
        )
        self.manuscript_id = create_manuscript(
            self.project_id,
            "Stability manuscript",
        )
        self.item_id = create_library_item(
            title="Stability of CM-01",
            authors="Jane Smith; Alex Researcher",
            publication_year=2025,
            source_name="Journal of Stability",
            doi="10.1000/stability",
            abstract="CM-01 remained stable under the selected conditions.",
            project_ids=[self.project_id],
        )

    def tearDown(self):
        if self.previous_db_path is None:
            os.environ.pop("DATABASE_PATH", None)
        else:
            os.environ["DATABASE_PATH"] = self.previous_db_path

        if self.previous_asset_path is None:
            os.environ.pop("MANUSCRIPT_ASSET_STORAGE_PATH", None)
        else:
            os.environ["MANUSCRIPT_ASSET_STORAGE_PATH"] = self.previous_asset_path

        self.temp_dir.cleanup()

    def _results_section(self):
        return next(
            section
            for section in get_manuscript_sections(self.manuscript_id)
            if section["section_type"] == "results"
        )

    def test_manuscript_outline_sources_and_citations(self):
        self.assertEqual(len(get_manuscript_sections(self.manuscript_id)), 7)
        self.assertEqual(len(get_project_library_sources(self.project_id)), 1)
        citation_key = attach_manuscript_source(self.manuscript_id, self.item_id)
        self.assertEqual(citation_key, "smith2025")

        section = self._results_section()
        update_manuscript_section(section["id"], content_md="A supported result.")
        token = insert_section_citation(section["id"], self.item_id)
        self.assertEqual(token, "[@smith2025]")
        self.assertIn(token, get_manuscript_section(section["id"])["content_md"])

        update_manuscript_source(
            self.manuscript_id,
            self.item_id,
            citation_key="stability2025",
            notes="Core reference",
        )
        updated = get_manuscript_section(section["id"])
        self.assertIn("[@stability2025]", updated["content_md"])
        self.assertNotIn("[@smith2025]", updated["content_md"])

    def test_versions_restore_and_duplicate_complete_snapshot(self):
        section = self._results_section()
        update_manuscript_section(section["id"], content_md="Original result.")
        attach_manuscript_source(self.manuscript_id, self.item_id)
        attach_manuscript_evidence(
            self.manuscript_id,
            "key_idea",
            999,
            "Stability window",
            "Near-neutral stability was strongest.",
        )
        version_id = create_manuscript_version(
            self.manuscript_id,
            "Original draft",
        )
        update_manuscript_section(section["id"], content_md="Changed result.")

        restore_manuscript_version(version_id)
        restored_section = next(
            row
            for row in get_manuscript_sections(self.manuscript_id)
            if row["section_type"] == "results"
        )
        self.assertEqual(restored_section["content_md"], "Original result.")
        self.assertEqual(len(get_manuscript_sources(self.manuscript_id)), 1)
        self.assertEqual(len(get_manuscript_evidence(self.manuscript_id)), 1)

        duplicate_id = duplicate_manuscript_version(version_id, "Duplicate draft")
        self.assertEqual(get_manuscript(duplicate_id)["title"], "Duplicate draft")
        self.assertEqual(len(get_manuscript_sections(duplicate_id)), 7)
        self.assertEqual(len(get_manuscript_versions(self.manuscript_id)), 1)

    def test_ai_response_and_all_export_formats(self):
        source_key = attach_manuscript_source(self.manuscript_id, self.item_id)
        self.assertEqual(source_key, "smith2025")
        section = self._results_section()
        update_manuscript_section(section["id"], content_md="Current result.")
        manuscript = get_manuscript(self.manuscript_id)
        sections = get_manuscript_sections(self.manuscript_id)
        sources = get_manuscript_sources(self.manuscript_id)
        result = generate_writing_suggestion(
            mode="Cite",
            instruction="Cite the result.",
            manuscript=manuscript,
            section=section,
            sections=sections,
            sources=sources,
            evidence=[],
            ai_provider=FixedWritingProvider(),
        )

        self.assertEqual(result["claims"][0]["status"], "supported")
        markdown = manuscript_markdown(manuscript, sections, sources)
        docx = manuscript_docx(manuscript, sections, sources)
        pdf = manuscript_pdf(manuscript, sections, sources)
        self.assertIn("# Stability manuscript", markdown)
        self.assertTrue(docx.startswith(b"PK"))
        self.assertTrue(pdf.startswith(b"%PDF"))

    def test_custom_ai_context_is_persisted_and_used(self):
        attach_manuscript_source(self.manuscript_id, self.item_id)
        attach_manuscript_evidence(
            self.manuscript_id,
            "key_idea",
            999,
            "Stability window",
            "Near-neutral stability was strongest.",
        )
        section = self._results_section()
        update_manuscript_section(section["id"], content_md="Selected context result.")
        saved = update_manuscript_ai_context(
            self.manuscript_id,
            context_mode="Custom",
            section_ids=[section["id"], 999999],
            source_ids=[self.item_id, 999999],
            evidence_keys=["key_idea:999", "summary:999999"],
        )
        self.assertEqual(saved["context_mode"], "Custom")
        self.assertEqual(saved["section_ids"], [section["id"]])
        self.assertEqual(saved["source_ids"], [self.item_id])
        self.assertEqual(saved["evidence_keys"], ["key_idea:999"])
        self.assertEqual(get_manuscript_ai_context(self.manuscript_id), saved)

        provider = FixedWritingProvider()
        generate_writing_suggestion(
            mode="Draft",
            instruction="Use only the custom context.",
            manuscript=get_manuscript(self.manuscript_id),
            section=get_manuscript_section(section["id"]),
            sections=get_manuscript_sections(self.manuscript_id),
            sources=get_manuscript_sources(self.manuscript_id),
            evidence=get_manuscript_evidence(self.manuscript_id),
            context_mode="Custom",
            context_sections=[get_manuscript_section(section["id"])],
            ai_provider=provider,
        )
        self.assertIn("CONTEXT_SCOPE: Custom", provider.prompt)
        self.assertIn("Selected context result.", provider.prompt)

    def test_multiple_citation_insertion_and_validation(self):
        second_item_id = create_library_item(
            title="Complementary stability evidence",
            authors="John Doe",
            publication_year=2024,
            source_name="Evidence Journal",
            doi="10.1000/complementary",
            abstract="A complementary stability analysis.",
            project_ids=[self.project_id],
        )
        attach_manuscript_source(self.manuscript_id, self.item_id)
        attach_manuscript_source(self.manuscript_id, second_item_id)
        section = self._results_section()
        update_manuscript_section(
            section["id"],
            content_md="First paragraph.\n\nSecond paragraph.",
        )
        tokens = insert_section_citations(
            section["id"],
            [self.item_id, second_item_id],
            placement="after_paragraph",
            after_paragraph=0,
        )
        updated = get_manuscript_section(section["id"])
        paragraphs = updated["content_md"].split("\n\n")
        self.assertEqual(paragraphs[1], " ".join(tokens))
        validation = validate_section_citations(
            updated["content_md"] + " [@missing2026]",
            get_manuscript_sources(self.manuscript_id),
        )
        self.assertEqual(len(validation["valid_keys"]), 2)
        self.assertEqual(validation["unknown_keys"], ["missing2026"])

    def test_scientific_objects_are_numbered_edited_and_referenced(self):
        results = self._results_section()
        introduction = next(
            section
            for section in get_manuscript_sections(self.manuscript_id)
            if section["section_type"] == "introduction"
        )
        later_figure_id = create_manuscript_asset(
            self.manuscript_id,
            results["id"],
            "figure",
            "Observed stability profile",
            storage_path=str(Path(self.temp_dir.name) / "manuscript-assets" / "missing.png"),
            mime_type="image/png",
        )
        first_figure_id = create_manuscript_asset(
            self.manuscript_id,
            introduction["id"],
            "figure",
            "Study workflow",
            storage_path=str(Path(self.temp_dir.name) / "manuscript-assets" / "workflow.png"),
            mime_type="image/png",
        )
        table_id = create_manuscript_asset(
            self.manuscript_id,
            results["id"],
            "table",
            "Stability measurements",
            content={
                "columns": ["Condition", "Value"],
                "rows": [{"Condition": "Room temperature", "Value": "98%"}],
            },
        )
        equation_id = create_manuscript_asset(
            self.manuscript_id,
            results["id"],
            "equation",
            "First-order rate equation",
            content={"latex": r"C_t = C_0 e^{-kt}"},
        )
        assets = get_manuscript_assets(self.manuscript_id)
        by_id = {asset["id"]: asset for asset in assets}
        self.assertEqual(by_id[first_figure_id]["label"], "Figure 1")
        self.assertEqual(by_id[later_figure_id]["label"], "Figure 2")
        self.assertEqual(by_id[table_id]["label"], "Table 1")
        self.assertEqual(by_id[equation_id]["label"], "Equation 1")

        update_manuscript_section(results["id"], content_md="The measurements are summarized in")
        token = insert_manuscript_asset_reference(results["id"], table_id)
        self.assertEqual(token, f"[[table:{table_id}]]")
        rendered = render_asset_references(
            get_manuscript_section(results["id"])["content_md"],
            assets,
        )
        self.assertEqual(rendered, "The measurements are summarized in Table 1")

        update_manuscript_asset(
            table_id,
            caption="Edited measurements",
            content={
                "columns": ["Condition", "Value", "Unit"],
                "rows": [{"Condition": "25 C", "Value": "98", "Unit": "%"}],
            },
        )
        updated = next(
            asset for asset in get_manuscript_assets(self.manuscript_id)
            if asset["id"] == table_id
        )
        self.assertEqual(updated["caption"], "Edited measurements")
        self.assertEqual(updated["content"]["columns"], ["Condition", "Value", "Unit"])
        move_manuscript_asset(equation_id, -1)
        ordered_results = get_manuscript_assets(
            self.manuscript_id,
            section_id=results["id"],
        )
        self.assertLess(
            [asset["id"] for asset in ordered_results].index(equation_id),
            [asset["id"] for asset in ordered_results].index(table_id),
        )

    def test_figure_upload_is_validated_and_normalized(self):
        image_bytes = BytesIO()
        Image.new("RGB", (320, 180), "#1473e6").save(image_bytes, format="JPEG")

        class Upload:
            name = "result figure.jpg"
            type = "image/jpeg"

            def getvalue(self):
                return image_bytes.getvalue()

        stored = save_figure_upload(Upload(), self.manuscript_id)
        self.assertEqual(stored["mime_type"], "image/png")
        self.assertEqual(stored["content"]["width"], 320)
        self.assertTrue(stored["storage_path"].endswith(".png"))
        self.assertTrue(read_manuscript_asset_file(stored["storage_path"]).startswith(b"\x89PNG"))

    def test_versions_remap_scientific_object_references(self):
        section = self._results_section()
        table_id = create_manuscript_asset(
            self.manuscript_id,
            section["id"],
            "table",
            "Versioned measurements",
            content={"columns": ["A"], "rows": [{"A": "1"}]},
        )
        insert_manuscript_asset_reference(section["id"], table_id)
        version_id = create_manuscript_version(self.manuscript_id, "With table")
        restore_manuscript_version(version_id)
        restored_assets = get_manuscript_assets(self.manuscript_id)
        self.assertEqual(len(restored_assets), 1)
        restored_section = next(
            row
            for row in get_manuscript_sections(self.manuscript_id)
            if row["section_type"] == "results"
        )
        self.assertIn(
            restored_assets[0]["reference_token"],
            restored_section["content_md"],
        )

        duplicate_id = duplicate_manuscript_version(version_id, "Objects duplicate")
        duplicate_assets = get_manuscript_assets(duplicate_id)
        duplicate_section = next(
            row
            for row in get_manuscript_sections(duplicate_id)
            if row["section_type"] == "results"
        )
        self.assertEqual(len(duplicate_assets), 1)
        self.assertIn(duplicate_assets[0]["reference_token"], duplicate_section["content_md"])

    def test_deleting_object_section_removes_cross_section_references(self):
        appendix_id = add_manuscript_section(self.manuscript_id, "Supplementary analysis")
        table_id = create_manuscript_asset(
            self.manuscript_id,
            appendix_id,
            "table",
            "Supplementary measurements",
            content={"columns": ["Value"], "rows": [{"Value": "42"}]},
        )
        results = self._results_section()
        insert_manuscript_asset_reference(results["id"], table_id)
        self.assertIn(
            f"[[table:{table_id}]]",
            get_manuscript_section(results["id"])["content_md"],
        )
        delete_manuscript_section(appendix_id)
        self.assertNotIn(
            f"[[table:{table_id}]]",
            get_manuscript_section(results["id"])["content_md"],
        )

    def test_scientific_objects_export_to_markdown_docx_and_pdf(self):
        section = self._results_section()
        figure_dir = Path(os.environ["MANUSCRIPT_ASSET_STORAGE_PATH"]) / str(self.manuscript_id)
        figure_dir.mkdir(parents=True, exist_ok=True)
        figure_path = figure_dir / "export.png"
        Image.new("RGB", (900, 420), "#dbeafe").save(figure_path)
        figure_id = create_manuscript_asset(
            self.manuscript_id,
            section["id"],
            "figure",
            "Concentration profile over time",
            alt_text="Blue concentration profile",
            storage_path=str(figure_path),
            mime_type="image/png",
            content={"width": 900, "height": 420},
        )
        table_id = create_manuscript_asset(
            self.manuscript_id,
            section["id"],
            "table",
            "Measured concentrations",
            content={
                "columns": ["Time", "Concentration"],
                "rows": [{"Time": "0 h", "Concentration": "100%"}],
            },
        )
        equation_id = create_manuscript_asset(
            self.manuscript_id,
            section["id"],
            "equation",
            "First-order degradation model",
            content={"latex": r"C_t = C_0 e^{-kt}"},
        )
        update_manuscript_section(
            section["id"],
            content_md=(
                f"The profile is shown in [[figure:{figure_id}]]. "
                f"Values appear in [[table:{table_id}]] and the model in "
                f"[[equation:{equation_id}]]."
            ),
        )
        manuscript = get_manuscript(self.manuscript_id)
        sections = get_manuscript_sections(self.manuscript_id)
        sources = get_manuscript_sources(self.manuscript_id)
        assets = get_manuscript_assets(self.manuscript_id)
        markdown = manuscript_markdown(manuscript, sections, sources, assets)
        docx_bytes = manuscript_docx(manuscript, sections, sources, assets)
        pdf_bytes = manuscript_pdf(manuscript, sections, sources, assets)
        self.assertIn("Figure 1", markdown)
        self.assertIn("Table 1", markdown)
        self.assertIn("Equation 1", markdown)
        self.assertIn("| Time | Concentration |", markdown)
        self.assertTrue(pdf_bytes.startswith(b"%PDF"))

        with zipfile.ZipFile(BytesIO(docx_bytes)) as archive:
            self.assertTrue(any(name.startswith("word/media/") for name in archive.namelist()))
            document_xml = archive.read("word/document.xml").decode("utf-8")
            self.assertIn("Figure 1", document_xml)
            self.assertIn("Table 1", document_xml)
            self.assertIn("Equation 1", document_xml)
            self.assertIn("Measured concentrations", document_xml)

    def test_text_formatting_is_preserved_in_docx_and_pdf(self):
        from docx import Document

        section = self._results_section()
        formatted_content = (
            "## pH-dependent stability\n"
            "The **compound** remained _highly stable_ according to the "
            "[protocol](https://example.com/protocol).\n\n"
            "- Neutral conditions\n"
            "- Low variability\n\n"
            "1. Prepare samples\n"
            "2. Run HPLC analysis\n\n"
            "> Stability remained above the acceptance threshold.\n\n"
            "Use `HPLC` for quantification."
        )
        update_manuscript_section(section["id"], content_md=formatted_content)
        manuscript = get_manuscript(self.manuscript_id)
        sections = get_manuscript_sections(self.manuscript_id)
        markdown = manuscript_markdown(manuscript, sections, [])
        docx_bytes = manuscript_docx(manuscript, sections, [])
        pdf_bytes = manuscript_pdf(manuscript, sections, [])
        self.assertIn("**compound**", markdown)
        self.assertTrue(pdf_bytes.startswith(b"%PDF"))

        document = Document(BytesIO(docx_bytes))
        heading = next(
            paragraph
            for paragraph in document.paragraphs
            if paragraph.text == "pH-dependent stability"
        )
        formatted_paragraph = next(
            paragraph
            for paragraph in document.paragraphs
            if "compound" in paragraph.text
        )
        self.assertEqual(heading.style.name, "Heading 2")
        self.assertTrue(any(run.text == "compound" and run.bold for run in formatted_paragraph.runs))
        self.assertTrue(any(run.text == "highly stable" and run.italic for run in formatted_paragraph.runs))
        self.assertEqual(
            len([p for p in document.paragraphs if p.style.name == "List Bullet"]),
            2,
        )
        self.assertEqual(
            len([p for p in document.paragraphs if p.style.name == "List Number"]),
            2,
        )

        with zipfile.ZipFile(BytesIO(docx_bytes)) as archive:
            document_xml = archive.read("word/document.xml").decode("utf-8")
            relationships = archive.read(
                "word/_rels/document.xml.rels"
            ).decode("utf-8")
            self.assertIn("w:hyperlink", document_xml)
            self.assertIn("https://example.com/protocol", relationships)

    def test_submission_profile_is_versioned_and_exported_to_docx(self):
        from docx import Document

        abstract = next(
            section
            for section in get_manuscript_sections(self.manuscript_id)
            if section["section_type"] == "abstract"
        )
        update_manuscript_section(
            abstract["id"],
            content_md="A complete abstract describing the objective, methods, results, and conclusion. " * 5,
        )
        profile = update_manuscript_submission_profile(
            self.manuscript_id,
            target_journal="Journal of Reproducible Research",
            journal_template="General IMRaD",
            short_title="CM-01 stability",
            authors=[{
                "name": "Jane Smith",
                "affiliations": "1",
                "email": "jane@example.com",
                "orcid": "0000-0001-2345-6789",
            }],
            affiliations=[{
                "label": "1",
                "institution": "Research Institute",
                "location": "Bucharest, Romania",
            }],
            corresponding_author={"name": "Jane Smith", "email": "jane@example.com"},
            keywords=["stability", "HPLC", "formulation"],
            word_limit=4200,
            abstract_word_limit=220,
            checklist={"author_approval": True},
        )
        version_id = create_manuscript_version(
            self.manuscript_id,
            "Before supervisor review",
            note="Ready for structured feedback.",
        )
        update_manuscript_submission_profile(
            self.manuscript_id,
            target_journal="Changed journal",
        )
        restore_manuscript_version(version_id)
        restored_profile = get_manuscript_submission_profile(self.manuscript_id)
        self.assertEqual(restored_profile["target_journal"], "Journal of Reproducible Research")
        self.assertEqual(restored_profile["authors"][0]["name"], "Jane Smith")

        update_manuscript_version(
            version_id,
            label="Before supervisor review",
            note="Supervisor comments requested on Results and Discussion.",
        )
        self.assertIn("Results", get_manuscript_version(version_id)["note"])
        add_manuscript_version_comment(
            version_id,
            "Please clarify the primary endpoint.",
            author_name="Supervisor",
        )
        comments = get_manuscript_version_comments(version_id)
        self.assertEqual(len(comments), 1)
        self.assertEqual(comments[0]["author_name"], "Supervisor")

        manuscript = get_manuscript(self.manuscript_id)
        sections = get_manuscript_sections(self.manuscript_id)
        docx_bytes = manuscript_docx(
            manuscript,
            sections,
            [],
            [],
            restored_profile,
        )
        document = Document(BytesIO(docx_bytes))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        self.assertIn("Jane Smith [1]", text)
        self.assertIn("Research Institute, Bucharest, Romania", text)
        self.assertIn("Corresponding author: Jane Smith · jane@example.com", text)
        self.assertIn("Keywords: stability, HPLC, formulation", text)

        with zipfile.ZipFile(BytesIO(docx_bytes)) as archive:
            footer_xml = archive.read("word/footer1.xml").decode("utf-8")
            self.assertIn("PAGE", footer_xml)

    def test_single_section_restore_preserves_other_current_sections(self):
        results = self._results_section()
        introduction = next(
            section
            for section in get_manuscript_sections(self.manuscript_id)
            if section["section_type"] == "introduction"
        )
        update_manuscript_section(results["id"], content_md="Versioned result.")
        update_manuscript_section(introduction["id"], content_md="Versioned introduction.")
        version_id = create_manuscript_version(self.manuscript_id, "Section baseline")
        update_manuscript_section(results["id"], content_md="Current changed result.")
        update_manuscript_section(introduction["id"], content_md="Keep this current introduction.")
        restore_manuscript_section_from_version(
            version_id,
            results["id"],
            results["id"],
        )
        self.assertEqual(get_manuscript_section(results["id"])["content_md"], "Versioned result.")
        self.assertEqual(
            get_manuscript_section(introduction["id"])["content_md"],
            "Keep this current introduction.",
        )

    def test_manuscript_checks_cover_evidence_abbreviations_structure_and_limits(self):
        attach_manuscript_source(self.manuscript_id, self.item_id)
        results = self._results_section()
        discussion = next(
            section
            for section in get_manuscript_sections(self.manuscript_id)
            if section["section_type"] == "discussion"
        )
        update_manuscript_section(
            results["id"],
            content_md=(
                "HPLC showed a significant increase of 42% under acidic conditions. "
                "The repeated finding remained clinically important."
            ),
        )
        update_manuscript_section(
            discussion["id"],
            content_md="A different mechanism decreased solubility in the final formulation.",
        )
        profile = update_manuscript_submission_profile(
            self.manuscript_id,
            word_limit=10,
            abstract_word_limit=5,
        )
        manuscript = get_manuscript(self.manuscript_id)
        sections = get_manuscript_sections(self.manuscript_id)
        sources = get_manuscript_sources(self.manuscript_id)
        checks = run_manuscript_checks(manuscript, sections, sources, profile)
        categories = {issue["category"] for issue in checks["issues"]}
        self.assertIn("Claims without evidence", categories)
        self.assertIn("Undefined abbreviations", categories)
        self.assertIn("Citations", categories)
        self.assertIn("Incomplete sections", categories)
        self.assertIn("Results vs Discussion", categories)
        self.assertIn("Word limits", categories)
        readiness = publication_readiness(
            manuscript,
            sections,
            sources,
            [],
            profile,
            checks,
        )
        self.assertFalse(readiness["ready"])
        self.assertLess(readiness["percent"], 100)


if __name__ == "__main__":
    unittest.main()
